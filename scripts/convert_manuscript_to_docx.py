#!/usr/bin/env python3
"""
Convert manuscript_combined.txt to a formatted Word document (.docx)
for Cell Systems Editorial Manager submission.

Formatting:
- 12pt Arial body text, justified
- 16pt bold Heading 1, 14pt bold Heading 2
- Double-spaced (line spacing = 2.0)
- 1-inch margins on all sides
- Page numbers in footer
- Superscript citations (Unicode superscript digits -> Word superscript runs)
- Greek letters and special characters preserved
- Key Resources Table and other tables rendered as Word tables
- Line numbers noted as manual step (python-docx has limited support)
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------- Constants ----------
FONT_NAME = "Arial"
BODY_SIZE = Pt(12)
HEADING_SIZE = Pt(16)
SUBHEADING_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)
TABLE_SIZE = Pt(9)

# Unicode superscript digit map
SUPERSCRIPT_MAP = {
    "\u2070": "0", "\u00b9": "1", "\u00b2": "2", "\u00b3": "3",
    "\u2074": "4", "\u2075": "5", "\u2076": "6", "\u2077": "7",
    "\u2078": "8", "\u2079": "9", "\u207b": "\u2212",  # superscript minus -> minus sign
    "\u207a": "+",  # superscript plus
}
SUPERSCRIPT_CHARS = set(SUPERSCRIPT_MAP.keys())

# Patterns for superscript sequences (citation numbers, exponents like 10⁻⁶)
SUPERSCRIPT_RE = re.compile(
    r"[\u2070\u00b9\u00b2\u00b3\u2074-\u2079\u207a\u207b]+"
)

# Subscript markup: `_{X}` (literal underscore, `{`, content with no `}` inside,
# `}`) renders X as a subscript run; the underscore and braces are consumed. A
# standalone `{` or `}` not preceded by `_` is left untouched (emitted literally).
SUBSCRIPT_MARKUP_RE = re.compile(r"_\{([^}]*)\}")

# Species binomials rendered in italic wherever they appear in body text.
# Word-boundary matched so partial substrings are never italicized; the source
# text is left unchanged (only the run's formatting differs).
SPECIES_RE = re.compile(r"\b(Homo sapiens|Mus musculus|Microcebus murinus)\b")


def decode_superscript(s: str) -> str:
    """Convert Unicode superscript characters to their normal equivalents."""
    return "".join(SUPERSCRIPT_MAP.get(c, c) for c in s)


def add_formatted_text(paragraph, text, font_name=FONT_NAME, font_size=BODY_SIZE,
                       bold=False, italic=False, superscript=False, subscript=False):
    """Add a run with specific formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript
    run.font.subscript = subscript
    # Set East Asia and complex script fonts too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
    return run


def add_text_with_superscripts(paragraph, text, font_name=FONT_NAME, font_size=BODY_SIZE,
                               bold=False, italic=False):
    """Add text to a paragraph, converting Unicode superscript chars to Word
    superscript runs and `_{...}` markup to subscript runs. Subscript markup,
    Unicode superscripts, and normal text interleave in left-to-right source
    order. Text containing no `_{...}` markup renders exactly as before."""
    # Italicize species binomials first: split on SPECIES_RE (capturing), emit
    # each matched binomial as an italic run (inheriting bold), and recurse on
    # the surrounding non-species text, which contains no binomial (so the
    # recursion terminates at depth 1 and behavior is unchanged for text with
    # no species).
    species_parts = SPECIES_RE.split(text)
    if len(species_parts) > 1:
        for sp_i, sp_seg in enumerate(species_parts):
            if not sp_seg:
                continue
            if sp_i % 2 == 1:
                add_formatted_text(paragraph, sp_seg, font_name, font_size, bold, True)
            else:
                add_text_with_superscripts(paragraph, sp_seg, font_name, font_size, bold, italic)
        return

    # Split on `_{...}` subscript markup. With one capturing group, re.split
    # yields alternating chunks: even index = literal text (handled by the
    # existing Unicode-superscript path below), odd index = the captured
    # subscript body (the `_{` and `}` are consumed). When the text contains no
    # markup this is a single literal chunk, so behavior is unchanged.
    segments = SUBSCRIPT_MARKUP_RE.split(text)
    for seg_i, seg in enumerate(segments):
        if seg_i % 2 == 1:
            # `_{X}` body -> one subscript run, same size as surrounding text.
            if seg:
                add_formatted_text(paragraph, seg, font_name, font_size,
                                   bold, italic, subscript=True)
            continue

        # Literal text: existing Unicode-superscript handling, unchanged.
        parts = SUPERSCRIPT_RE.split(seg)
        sups = SUPERSCRIPT_RE.findall(seg)

        for i, part in enumerate(parts):
            if part:
                add_formatted_text(paragraph, part, font_name, font_size, bold, italic)
            if i < len(sups):
                decoded = decode_superscript(sups[i])
                add_formatted_text(paragraph, decoded, font_name, font_size, bold, italic, superscript=True)


def set_paragraph_spacing(paragraph, line_spacing=2.0, space_after=Pt(0), space_before=Pt(0)):
    """Set paragraph line spacing and before/after spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before


def set_paragraph_alignment(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Set paragraph alignment."""
    paragraph.alignment = alignment


def add_page_number(doc):
    """Add page numbers to the footer of all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add PAGE field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        for r in [run, run2, run3]:
            r.font.name = FONT_NAME
            r.font.size = Pt(10)


def add_line_numbers(doc):
    """
    Add continuous line numbers via section properties.
    python-docx doesn't natively support this, so we use raw XML.
    """
    for section in doc.sections:
        sectPr = section._sectPr
        ln_num = parse_xml(
            f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:restart="continuous"/>'
        )
        sectPr.append(ln_num)


def parse_sections(text: str) -> list:
    """
    Parse the manuscript text into sections delimited by ======== lines.
    Returns list of (section_title, section_body) tuples.

    Structure pattern:
        ========================================================================
        SECTION TITLE
        ========================================================================

        section body content...
    """
    # Use regex to find section blocks: delimiter, title, delimiter, body
    delimiter_pattern = r"={10,}"  # 10+ equals signs
    # Split into blocks by delimiter lines
    lines = text.split("\n")
    sections = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Look for delimiter line
        if re.match(delimiter_pattern, line):
            # Next non-empty line should be the title
            i += 1
            title = ""
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                title = lines[i].strip()
                i += 1
            # Skip closing delimiter
            while i < len(lines):
                if re.match(delimiter_pattern, lines[i].strip()):
                    i += 1
                    break
                elif not lines[i].strip():
                    i += 1
                else:
                    break

            # Collect body until next delimiter
            body_lines = []
            while i < len(lines) and not re.match(delimiter_pattern, lines[i].strip()):
                body_lines.append(lines[i])
                i += 1

            body = "\n".join(body_lines).strip()
            sections.append((title, body))
        else:
            i += 1

    return sections


def is_table_block(lines: list) -> bool:
    """Check if a block of lines forms a markdown table."""
    if len(lines) < 2:
        return False
    pipe_lines = [l for l in lines if "|" in l]
    return len(pipe_lines) >= 2


SEP_ROW = re.compile(r'^\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$')


def _pipe_run_has_separator(lines, start):
    """Return True if the contiguous run of pipe-containing lines starting at
    index `start` includes a markdown table separator row. Distinguishes a
    genuine table (which has a separator) from prose containing '|'."""
    k = start
    while k < len(lines):
        lk = lines[k].strip()
        if not lk or "|" not in lk:
            break
        if SEP_ROW.match(lk):
            return True
        k += 1
    return False


# STAR Methods run-in lead-in labels: a short bold lead phrase ending in a
# period, followed by body text on the same line. Matched by prefix in
# process_body_content and rendered as a bold-led body paragraph (the remainder
# after the lead stays non-bold within the same paragraph).
STAR_RUNIN = {
    'Primary dataset.',
    'Sun2023[24].',
    'PanSci[27].',
    'CellHint[28].',
    'Cell type matching.',
    'Preprocessing.',
    'Ortholog space.',
    'Primary analysis.',
    'Sensitivity analysis (no-immune subset).',
    'Per-type residual ranking.',
    'Layer 1: Centroid position.',
    'Layer 2: Covariance ellipsoid orientation.',
    'Distinctness of centroid and ellipsoid conservation.',
    'Conserved principal component genes.',
    'Eigenvalue profile similarity (exploratory; not carried as an inferential finding).',
    'Sensitivity to ribosomal-protein-gene contribution.',
    'Independent replication on the Human Protein Atlas.',
    'Overall tree structure.',
    'Anticorrelation observation (post-hoc/exploratory).',
    'Mechanistic hypotheses ruled out.',
    'Key software packages.',
    'Primary datasets.',
    'Ortholog reference.',
}


def parse_table(lines: list) -> list:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Split on pipe
        cells = [c.strip() for c in line.split("|")]
        # Remove empty first/last from leading/trailing pipes
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        # Check if this is a separator line (all dashes/colons)
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        if cells:
            rows.append(cells)
    return rows


def add_word_table(doc, rows, bold_header=True, col_widths=None):
    """Add a formatted Word table from parsed rows."""
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = row.cells[j]
                # Clear default paragraph
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                is_bold = bold_header and i == 0

                # Check for bold markers in cell text
                cell_text_clean = cell_text
                if cell_text.startswith("**") and cell_text.endswith("**"):
                    cell_text_clean = cell_text[2:-2]
                    is_bold = True
                elif cell_text.startswith("**"):
                    cell_text_clean = cell_text.replace("**", "")
                    is_bold = True

                add_text_with_superscripts(
                    p, cell_text_clean, font_size=TABLE_SIZE, bold=is_bold
                )
                set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(2), space_before=Pt(2))

    if col_widths:
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = col_widths[ci]
    return table


def _start_landscape(doc):
    from docx.enum.section import WD_SECTION
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = Inches(0.75)
    sec.top_margin = sec.bottom_margin = Inches(1)


def _end_landscape(doc):
    from docx.enum.section import WD_SECTION
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)


def process_paragraph_text(doc, text, font_size=BODY_SIZE, bold=False, italic=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=2.0,
                           space_after=Pt(0), space_before=Pt(0), first_line_indent=None):
    """Add a fully formatted paragraph with superscript handling."""
    p = doc.add_paragraph()
    add_text_with_superscripts(p, text, font_size=font_size, bold=bold, italic=italic)
    set_paragraph_spacing(p, line_spacing=line_spacing, space_after=space_after, space_before=space_before)
    set_paragraph_alignment(p, alignment)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def add_section_heading(doc, title, level=1):
    """Add a section heading using proper Word heading styles (Heading 1/2/3)."""
    if level == 1:
        p = doc.add_heading(title, level=1)
        # Override style font to match our spec
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                                   f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
                rPr.insert(0, rFonts)
            else:
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rFonts.set(qn(attr), FONT_NAME)
        set_paragraph_spacing(p, line_spacing=2.0, space_before=Pt(24), space_after=Pt(12))
    elif level == 2:
        p = doc.add_heading(title, level=2)
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = SUBHEADING_SIZE
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                                   f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
                rPr.insert(0, rFonts)
            else:
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rFonts.set(qn(attr), FONT_NAME)
        set_paragraph_spacing(p, line_spacing=2.0, space_before=Pt(18), space_after=Pt(6))
    else:
        p = doc.add_heading(title, level=3)
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = HEADING3_SIZE
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                                   f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
                rPr.insert(0, rFonts)
            else:
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rFonts.set(qn(attr), FONT_NAME)
        set_paragraph_spacing(p, line_spacing=2.0, space_before=Pt(12), space_after=Pt(6))
    set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.LEFT)
    return p


def process_body_content(doc, body: str, is_star_methods=False):
    """Process the body text of a section, handling paragraphs, tables, bullets, and sub-headings."""
    lines = body.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Handle horizontal rules (--- separators in supplemental legends)
        if line == "---":
            # Add a thin horizontal line or just skip
            i += 1
            continue

        # Detect table blocks (at least 2 consecutive pipe-containing lines)
        if line.count("|") >= 2 or ("|" in line and i + 1 < len(lines) and "|" in lines[i + 1]):
            table_lines = []
            j = i
            while j < len(lines):
                lj = lines[j].strip()
                if not lj:
                    break
                if "|" in lj or re.match(r"^[-:| ]+$", lj):
                    table_lines.append(lj)
                else:
                    break
                j += 1

            if len(table_lines) >= 2 and any(SEP_ROW.match(l) for l in table_lines):
                rows = parse_table(table_lines)
                if rows and len(rows) >= 2:
                    add_word_table(doc, rows)
                    # Add spacing after table
                    p = doc.add_paragraph()
                    set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0))
                    i = j
                    continue

        # Handle bullet points
        if line.startswith("- "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph()
            # Add bullet character
            add_formatted_text(p, "\u2022 ")
            add_text_with_superscripts(p, bullet_text)
            set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0))
            set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            i += 1
            continue

        # STAR Methods headings — known section and sub-section names
        if is_star_methods and not line.startswith("|"):
            # Level 2: Major STAR Methods sections (ALL CAPS)
            star_major = {
                "RESOURCE AVAILABILITY", "EXPERIMENTAL MODEL AND SUBJECT DETAILS",
                "METHOD DETAILS", "QUANTIFICATION AND STATISTICAL ANALYSIS",
                "ADDITIONAL RESOURCES",
            }
            if line in star_major:
                add_section_heading(doc, line, level=2)
                i += 1
                continue

            # Level 3: Known sub-section headings
            star_sub = {
                "Data acquisition and cell type selection",
                "Ortholog mapping and gene space", "Normalization and centroid computation",
                "PCA interdependence", "Independent PCA sensitivity analysis",
                "Procrustes superimposition", "Human-versus-human negative control",
                "Replication datasets", "Harmonized centroid replication",
                "Macaque extension", "Two-layer geometric conservation",
                "SAMap validation and comparison", "CellMarker identity gene set comparison",
                "L1000 landmark gene analysis", "Treeness analysis",
                "DILI analysis", "Cancer deformation analysis",
                "Permutation test and statistical framework",
                "Leave-one-out cross-validation", "Bootstrap robustness",
                "Mantel test", "Cell count confound analysis",
                "Mechanistic null tests", "Simulation study",
                "Bootstrap ranking analysis", "Expanded negative controls",
                "Biological predictors", "Mouse lemur analysis",
                "Pan-Census replication", "Software and reproducibility",
                "Use of generative AI",
                "Species-mixing concordance", "Unknown cell exclusion",
                "Primary three-species analysis",
                "Hepatocyte rigidity across three species",
                "Full rigidity ranking", "Conserved principal component genes",
                "Donor-split within-species control", "Replication diagnostics",
                "KEY RESOURCES TABLE",
                "Layer 2 PanSci replication",
            }
            if line in star_sub:
                add_section_heading(doc, line, level=2)
                i += 1
                continue

            # STAR Methods run-in lead-in: bold the lead phrase, keep the
            # remaining body text non-bold in the same paragraph. Rendered as a
            # single line so trailing body sentences stay their own paragraphs.
            runin = next((ph for ph in STAR_RUNIN if line.startswith(ph)), None)
            if runin:
                rest = line[len(runin):]
                p = doc.add_paragraph()
                add_text_with_superscripts(p, runin, font_size=BODY_SIZE, bold=True)
                if rest:
                    add_text_with_superscripts(p, rest, font_size=BODY_SIZE)
                set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0), space_before=Pt(0))
                set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
                i += 1
                continue

        # Regular paragraph - collect continuation lines
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            next_l = lines[j].strip()
            # Stop at empty line, bullet, table-like line, or separator
            if (not next_l or next_l.startswith("- ") or next_l == "---"
                    or ("|" in next_l and _pipe_run_has_separator(lines, j))
                    or (is_star_methods and any(next_l.startswith(ph) for ph in STAR_RUNIN))):
                break
            para_lines.append(next_l)
            j += 1

        full_para = " ".join(para_lines)

        # Check if this paragraph starts with bold markers (like **Figure S1...**)
        if full_para.startswith("**") and "**" in full_para[2:]:
            # Find the end of the bold section
            end_bold = full_para.index("**", 2)
            bold_text = full_para[2:end_bold]
            rest_text = full_para[end_bold + 2:]

            p = doc.add_paragraph()
            add_text_with_superscripts(p, bold_text, bold=True)
            if rest_text:
                add_text_with_superscripts(p, rest_text)
            set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0))
            set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        else:
            process_paragraph_text(doc, full_para)

        i = j


def build_document(input_path: str, output_path: str):
    """Build the Word document from the manuscript text."""

    # Read input
    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()

    # Parse sections
    sections = parse_sections(text)

    # Create document
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Set default font for Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                           f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)
        rFonts.set(qn("w:eastAsia"), FONT_NAME)
        rFonts.set(qn("w:cs"), FONT_NAME)

    # Configure Heading styles with Arial and proportional sizes
    for hlevel, hsize in [(1, HEADING_SIZE), (2, SUBHEADING_SIZE), (3, HEADING3_SIZE)]:
        hstyle = doc.styles[f"Heading {hlevel}"]
        hstyle.font.name = FONT_NAME
        hstyle.font.size = hsize
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        hPr = hstyle.element.get_or_add_rPr()
        hFonts = hPr.find(qn("w:rFonts"))
        if hFonts is None:
            hFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                               f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
            hPr.insert(0, hFonts)
        else:
            hFonts.set(qn("w:ascii"), FONT_NAME)
            hFonts.set(qn("w:hAnsi"), FONT_NAME)
            hFonts.set(qn("w:eastAsia"), FONT_NAME)
            hFonts.set(qn("w:cs"), FONT_NAME)

    # Process each section
    for sec_title, sec_body in sections:
        if not sec_body and not sec_title:
            continue

        sec_title_clean = sec_title.strip()

        # --- HIGHLIGHTS AND eTOC BLURB ---
        if "HIGHLIGHTS" in sec_title_clean.upper() and "ETOC" in sec_title_clean.upper():
            add_section_heading(doc, "HIGHLIGHTS AND eTOC BLURB")

            # Parse highlights and eTOC
            # Use last split — earlier parts may match "eTOC Blurb" in the NOTE line
            parts = sec_body.split("eTOC Blurb")
            if len(parts) >= 2:
                highlights_text = "eTOC Blurb".join(parts[:-1])
                etoc_text = parts[-1].strip()
                if etoc_text.startswith("(In Brief)"):
                    etoc_text = etoc_text[len("(In Brief)"):].strip()

                # Highlights sub-heading
                add_section_heading(doc, "Highlights", level=2)
                for line in highlights_text.split("\n"):
                    line = line.strip()
                    if line.startswith("- "):
                        bullet_text = line[2:].strip()
                        p = doc.add_paragraph()
                        add_formatted_text(p, "\u2022 ")
                        add_text_with_superscripts(p, bullet_text)
                        set_paragraph_spacing(p, line_spacing=2.0)
                        set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(-0.25)
                    elif line and line != "Highlights":
                        process_paragraph_text(doc, line)

                # eTOC Blurb sub-heading
                add_section_heading(doc, "In Brief", level=2)
                process_paragraph_text(doc, etoc_text)
            else:
                process_body_content(doc, sec_body)

        # --- TITLE PAGE ---
        elif sec_title_clean == "TITLE PAGE":
            # Title - centered, bold, 14pt
            lines = sec_body.split("\n")
            title_lines = []
            author_lines = []
            in_authors = False
            short_title = None

            for line in lines:
                line_s = line.strip()
                if not line_s:
                    if title_lines and not in_authors:
                        in_authors = True
                    continue
                if line_s.startswith("Short title:"):
                    short_title = line_s
                    continue
                if line_s.startswith("Authors:"):
                    in_authors = True
                    continue
                if in_authors:
                    author_lines.append(line_s)
                else:
                    title_lines.append(line_s)

            # Add title
            title_text = " ".join(title_lines)
            p = doc.add_paragraph()
            add_formatted_text(p, title_text, font_size=HEADING_SIZE, bold=True)
            set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(24))
            set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)

            # Short title (running head), if present
            if short_title:
                p = doc.add_paragraph()
                add_formatted_text(p, short_title, italic=True)
                set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(12))
                set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)

            # Add author info
            for aline in author_lines:
                p = doc.add_paragraph()
                add_text_with_superscripts(p, aline)
                set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0))
                set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)

        # --- Abstract (PLOS) / SUMMARY (legacy Cell) ---
        elif sec_title_clean in ("Abstract", "SUMMARY"):
            add_section_heading(doc, "Abstract" if sec_title_clean == "Abstract" else "Summary")

            # Split body into summary paragraph and keywords
            parts = sec_body.split("Keywords:")
            summary_text = parts[0].strip()
            process_paragraph_text(doc, summary_text)

            if len(parts) > 1:
                keywords = parts[1].strip()
                p = doc.add_paragraph()
                add_formatted_text(p, "Keywords: ", bold=True)
                add_text_with_superscripts(p, keywords)
                set_paragraph_spacing(p, line_spacing=2.0, space_before=Pt(12))
                set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # --- Standard text sections ---
        elif sec_title_clean in ("INTRODUCTION", "DISCUSSION"):
            add_section_heading(doc, sec_title_clean.title())
            process_body_content(doc, sec_body)

        # --- RESULTS (has sub-headings) ---
        elif sec_title_clean == "RESULTS":
            add_section_heading(doc, "Results")

            lines = sec_body.split("\n")
            i = 0
            landscape_open = False
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # Inline table block (e.g. Table 1) — mirror process_body_content
                if line.count("|") >= 2 or ("|" in line and i + 1 < len(lines) and "|" in lines[i + 1]):
                    table_lines = []
                    j = i
                    while j < len(lines):
                        lj = lines[j].strip()
                        if not lj:
                            break
                        if "|" in lj or re.match(r"^[-:| ]+$", lj):
                            table_lines.append(lj)
                        else:
                            break
                        j += 1
                    if len(table_lines) >= 2 and any(SEP_ROW.match(l) for l in table_lines):
                        rows = parse_table(table_lines)
                        if rows and len(rows) >= 2:
                            n_cols = max(len(r) for r in rows)
                            cw = None
                            if n_cols >= 7 and landscape_open:
                                W = [6, 18, 72, 24, 7, 11, 11, 13, 14, 27, 14]
                                if len(W) == n_cols:
                                    cw = [Inches(w / sum(W) * 9.5) for w in W]
                            add_word_table(doc, rows, col_widths=cw)
                            p = doc.add_paragraph()
                            set_paragraph_spacing(p, line_spacing=2.0, space_after=Pt(0))
                            i = j
                            continue

                next_line = ""
                for _k in range(i + 1, min(i + 4, len(lines))):
                    if lines[_k].strip():
                        next_line = lines[_k].strip()
                        break
                is_subheading = (
                    len(line) < 100
                    and not line.endswith(".")
                    and not line.startswith("-")
                    and not line.startswith("|")
                    and not line.startswith("(")
                    and (line[0].isupper() or re.match(r"^\d+\.\s+[A-Z]", line))
                    and next_line
                    and len(next_line) > len(line)
                )
                if is_subheading:
                    add_section_heading(doc, line, level=2)
                    i += 1
                    continue

                para_lines = [line]
                j = i + 1
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl or "|" in nl:
                        break
                    peek = lines[j + 1].strip() if j + 1 < len(lines) else ""
                    if (len(nl) < 100 and not nl.endswith(".")
                            and not nl.startswith("-") and not nl.startswith("|")
                            and not nl.startswith("(")
                            and (nl[0].isupper() or re.match(r"^\d+\.\s+[A-Z]", nl))
                            and peek and len(peek) > len(nl)):
                        break
                    para_lines.append(nl)
                    j += 1
                full_para = " ".join(para_lines)
                nxt = lines[j].strip() if j < len(lines) else ""
                if nxt and ("|" in nxt) and (nxt.count("|") + 1) >= 7:
                    _start_landscape(doc)
                    landscape_open = True
                    process_paragraph_text(doc, full_para)
                else:
                    process_paragraph_text(doc, full_para)
                    if landscape_open:
                        _end_landscape(doc)
                        landscape_open = False
                i = j
            if landscape_open:
                _end_landscape(doc)

        # --- ACKNOWLEDGMENTS, AUTHOR CONTRIBUTIONS ---
        # Cell-era declaration special-cases removed in the PLOS reformat: generative-AI / legacy "AI DISCLOSURE" (now the Methods "Use of generative AI" subsection) and "DECLARATION OF INTERESTS" (now PLOS "Competing Interests"). ACKNOWLEDGMENTS / AUTHOR CONTRIBUTIONS keep their explicit branch above.
        elif sec_title_clean in ("ACKNOWLEDGMENTS", "AUTHOR CONTRIBUTIONS"):
            add_section_heading(doc, sec_title_clean.title())
            process_body_content(doc, sec_body)

        # --- FIGURE LEGENDS ---
        elif sec_title_clean == "FIGURE LEGENDS":
            add_section_heading(doc, "Figure Legends")

            lines = sec_body.split("\n")
            i = 0
            # Skip "Figure Legends" duplicate title line
            if lines and lines[0].strip() == "Figure Legends":
                i = 1

            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # Collect this legend: title line plus its body lines, breaking
                # at the next "Figure N."/"Table N." boundary or a blank line so
                # each legend becomes its own paragraph (the .txt has no
                # blank/--- separators between main legends).
                para_lines = [line]
                j = i + 1
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl:
                        break
                    if re.match(r"^(Figure \d+\.|Table \d+\.)", nl):
                        break
                    para_lines.append(nl)
                    j += 1

                # Bold the title (ref + title); body (panel descriptions)
                # non-bold. One paragraph per legend, matching
                # supplementary-legend rendering.
                title_line = para_lines[0]
                title_match = re.match(r"^(Figure \d+\.|Table \d+\.)", title_line)
                if title_match:
                    p = doc.add_paragraph()
                    if len(para_lines) > 1:
                        # Multi-line legend (the six figure legends): line-based
                        # split — first physical line bold, body lines non-bold.
                        body_text = " ".join(para_lines[1:])
                        add_text_with_superscripts(p, title_line, bold=True)
                        if body_text:
                            add_text_with_superscripts(p, " " + body_text)
                    else:
                        # Single-line legend (Table 1): sentence-based split —
                        # bold the label plus its first sentence (through that
                        # sentence's terminal period), remainder non-bold, all in
                        # the one paragraph.
                        label = title_match.group(1)
                        sent_end = re.search(r"\.(?=\s|$)", title_line[len(label):])
                        if sent_end:
                            split = len(label) + sent_end.end()
                            add_text_with_superscripts(p, title_line[:split], bold=True)
                            if title_line[split:]:
                                add_text_with_superscripts(p, title_line[split:])
                        else:
                            add_text_with_superscripts(p, title_line, bold=True)
                    set_paragraph_spacing(p, line_spacing=2.0, space_before=Pt(12))
                    set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
                else:
                    process_paragraph_text(doc, " ".join(para_lines))

                i = j

        # --- Materials and Methods (PLOS) / STAR METHODS (legacy Cell) ---
        # is_star_methods drives the level-2/level-3 method sub-heading and
        # run-in detection; the PLOS rename must keep that flag set.
        elif sec_title_clean in ("Materials and Methods", "STAR METHODS"):
            heading = ("Materials and Methods"
                       if sec_title_clean == "Materials and Methods" else "STAR Methods")
            add_section_heading(doc, heading)
            process_body_content(doc, sec_body, is_star_methods=True)

        # --- KEY RESOURCES TABLE ---
        elif sec_title_clean == "KEY RESOURCES TABLE":
            # This is handled inline if it appears, but also might be in STAR Methods
            add_section_heading(doc, "Key Resources Table")

            # Parse the table from the body
            lines = sec_body.split("\n")
            table_lines = [l for l in lines if "|" in l]
            if table_lines:
                rows = parse_table(table_lines)
                if rows:
                    add_word_table(doc, rows)

        # --- Supporting Information (PLOS) / SUPPLEMENTAL (legacy Cell) ---
        elif "SUPPLEMENTAL" in sec_title_clean.upper() or "SUPPORTING" in sec_title_clean.upper():
            heading = ("Supporting Information"
                       if "SUPPORTING" in sec_title_clean.upper() else "Supplemental Information")
            add_section_heading(doc, heading)
            process_body_content(doc, sec_body)

        # --- REFERENCES ---
        elif sec_title_clean == "REFERENCES":
            add_section_heading(doc, "References")

            lines = sec_body.split("\n")
            i = 0
            # Skip "References" duplicate title line
            if lines and lines[0].strip() == "References":
                i = 1

            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # Collect multi-line reference
                ref_lines = [line]
                j = i + 1
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl:
                        break
                    # New reference starts with a number followed by period
                    if re.match(r"^\d+\.", nl):
                        break
                    ref_lines.append(nl)
                    j += 1

                full_ref = " ".join(ref_lines)
                process_paragraph_text(
                    doc, full_ref,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_after=Pt(6),
                )
                i = j

        else:
            # Generic section
            if sec_title_clean:
                add_section_heading(doc, sec_title_clean.title())
            process_body_content(doc, sec_body)

    # --- Post-processing ---

    # Add page numbers
    add_page_number(doc)

    # Add line numbers (via section properties XML)
    add_line_numbers(doc)

    # Neutralize document properties for the public/submission artifact: set a
    # single human author and clear the tool-generated description so no library
    # name, machine user, or institutional identity is stamped into the .docx
    # (and, downstream, the docx2pdf-converted .pdf). No employer / non-
    # Sriram identity is ever written.
    cp = doc.core_properties
    cp.author = "Sriram Devadas"
    cp.last_modified_by = "Sriram Devadas"
    cp.title = "Quantifying the Conserved Geometry of Cell-Type Identity Across Mammalian Species"
    cp.comments = ""
    cp.category = ""

    # Save
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"  - Font: {FONT_NAME}, body {BODY_SIZE}, H1 {HEADING_SIZE}, H2 {SUBHEADING_SIZE}")
    print(f"  - Margins: 1 inch all sides")
    print(f"  - Line spacing: double (2.0)")
    print(f"  - Page numbers: footer, centered")
    print(f"  - Line numbers: added via XML (verify in Word)")
    print(f"  - Superscript citations: converted from Unicode to Word superscript")
    print()
    print("MANUAL STEPS RECOMMENDED:")
    print("  1. Open in Word and verify line numbers display correctly")
    print("  2. Update page fields (Ctrl+A, then F9) to refresh page numbers")
    print("  3. Verify all superscript citations render correctly")
    print("  4. Check table formatting and adjust column widths if needed")


if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parent.parent
    input_file = repo_root / "docs" / "submission" / "manuscript_combined.txt"
    output_file = repo_root / "docs" / "submission" / "cellwarp_manuscript.docx"

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    build_document(str(input_file), str(output_file))
