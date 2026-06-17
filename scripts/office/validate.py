#!/usr/bin/env python3
"""
Validate a Cell Systems submission manuscript .docx file.

Checks:
- File opens as valid docx
- Expected sections present (via Heading 1/2 styles)
- Font is Arial
- Page setup: US Letter, 1-inch margins
- Double-spaced body text
- Superscript runs present (citation numbers)
- Greek/special characters preserved
- Line numbers enabled
- Tables present
- Page count estimate
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml.ns import qn


def validate(docx_path: str) -> dict:
    """Validate a manuscript docx and return a report dict."""
    path = Path(docx_path)
    if not path.exists():
        print(f"ERROR: File not found: {docx_path}")
        sys.exit(1)

    doc = Document(str(path))
    warnings = []
    info = {}

    # --- Section inventory ---
    headings = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = p.style.name.replace("Heading ", "")
            headings.append((level, p.text.strip()))

    h1_titles = [t for lvl, t in headings if lvl == "1"]
    info["heading_1_count"] = len(h1_titles)
    info["heading_2_count"] = len([t for lvl, t in headings if lvl == "2"])
    info["heading_3_count"] = len([t for lvl, t in headings if lvl == "3"])
    info["total_headings"] = len(headings)
    info["h1_sections"] = h1_titles

    # Expected PLOS top-level sections
    expected_h1 = [
        "Abstract", "Author Summary", "Introduction", "Results", "Discussion",
        "Materials and Methods", "Acknowledgments", "Statements & Declarations",
        "Funding", "Author Contributions", "Competing Interests",
        "Data Availability", "References", "Figure Legends",
        "Supporting Information",
    ]
    # Fuzzy match: check that each expected section is present (case-insensitive partial)
    found_sections = set()
    for exp in expected_h1:
        for title in h1_titles:
            if exp.lower() in title.lower():
                found_sections.add(exp)
                break
    missing = [s for s in expected_h1 if s not in found_sections]
    if missing:
        warnings.append(f"Missing expected H1 sections: {', '.join(missing)}")
    info["expected_sections_found"] = len(found_sections)
    info["expected_sections_total"] = len(expected_h1)

    # --- Font check ---
    font_names = set()
    normal_style = doc.styles["Normal"]
    if normal_style.font.name:
        font_names.add(normal_style.font.name)
    # Sample first 50 paragraphs for run-level fonts
    for p in doc.paragraphs[:50]:
        for run in p.runs:
            if run.font.name:
                font_names.add(run.font.name)
    info["fonts_detected"] = sorted(font_names)
    if "Arial" not in font_names:
        warnings.append("Arial font not detected in document")

    # --- Page setup ---
    section = doc.sections[0]
    page_w = section.page_width
    page_h = section.page_height
    info["page_width_in"] = round(page_w / 914400, 2)
    info["page_height_in"] = round(page_h / 914400, 2)
    # US Letter = 8.5 x 11
    if abs(page_w / 914400 - 8.5) > 0.1 or abs(page_h / 914400 - 11) > 0.1:
        warnings.append(f"Page size not US Letter: {info['page_width_in']}×{info['page_height_in']} in")

    # Margins (1 inch = 914400 EMU)
    margins = {
        "top": section.top_margin,
        "bottom": section.bottom_margin,
        "left": section.left_margin,
        "right": section.right_margin,
    }
    info["margins_in"] = {k: round(v / 914400, 2) for k, v in margins.items()}
    for name, val in margins.items():
        if abs(val / 914400 - 1.0) > 0.15:
            warnings.append(f"{name} margin is {round(val / 914400, 2)} in, expected ~1.0 in")

    # --- Line spacing ---
    double_spaced_count = 0
    total_body = 0
    for p in doc.paragraphs:
        if p.style.name == "Normal" and p.text.strip():
            total_body += 1
            if p.paragraph_format.line_spacing and abs(p.paragraph_format.line_spacing - 2.0) < 0.1:
                double_spaced_count += 1
    info["body_paragraphs"] = total_body
    info["double_spaced"] = double_spaced_count
    if total_body > 0 and double_spaced_count < total_body * 0.8:
        warnings.append(f"Only {double_spaced_count}/{total_body} body paragraphs are double-spaced")

    # --- Superscript check ---
    superscript_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.superscript:
                superscript_count += 1
    info["superscript_runs"] = superscript_count
    if superscript_count < 10:
        warnings.append(f"Only {superscript_count} superscript runs found (expected many citations)")

    # --- Special characters ---
    full_text = "\n".join(p.text for p in doc.paragraphs)
    greek_chars = set(re.findall(r"[ρα-ωΑ-Ω]", full_text))
    special_chars = set(re.findall(r"[≤≥×−±°μ]", full_text))
    info["greek_letters"] = sorted(greek_chars)
    info["special_characters"] = sorted(special_chars)
    if not greek_chars:
        warnings.append("No Greek letters found in text")
    if not special_chars:
        warnings.append("No special characters (≤, ×, −) found in text")

    # --- Line numbers ---
    has_line_nums = False
    for sec in doc.sections:
        sectPr = sec._sectPr
        ln = sectPr.find(qn("w:lnNumType"))
        if ln is not None:
            has_line_nums = True
            break
    info["line_numbers"] = has_line_nums
    if not has_line_nums:
        warnings.append("Line numbers not enabled in section properties")

    # --- Tables ---
    info["table_count"] = len(doc.tables)
    if len(doc.tables) == 0:
        warnings.append("No tables found")

    # --- Placeholder fields ---
    placeholders_found = []
    for pattern in [r"\[AUTHOR\]", r"\[EMAIL\]", r"\[REPOSITORY URL", r"\[Name\]", r"\[email"]:
        if re.search(pattern, full_text):
            placeholders_found.append(pattern.strip("\\").strip("[]"))
    info["placeholders_preserved"] = placeholders_found

    # --- Page estimate ---
    # Rough estimate: ~250 words/page double-spaced
    word_count = len(full_text.split())
    info["word_count"] = word_count
    info["estimated_pages"] = max(1, round(word_count / 250))

    # --- Total paragraph count ---
    info["total_paragraphs"] = len(doc.paragraphs)

    return info, warnings


def main():
    if len(sys.argv) < 2:
        print("Usage: python validate.py <manuscript.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    # Resolve relative to docs/submission/ if not absolute
    p = Path(docx_path)
    if not p.is_absolute() and not p.exists():
        alt = Path(__file__).resolve().parent.parent.parent / "docs" / "submission" / p.name
        if alt.exists():
            p = alt
    docx_path = str(p)

    info, warnings = validate(docx_path)

    print("=" * 60)
    print("MANUSCRIPT VALIDATION REPORT")
    print("=" * 60)
    print()
    print(f"File: {docx_path}")
    print(f"Word count: {info['word_count']:,}")
    print(f"Estimated pages: {info['estimated_pages']}")
    print(f"Total paragraphs: {info['total_paragraphs']}")
    print()
    print("--- Page Setup ---")
    print(f"Page size: {info['page_width_in']} × {info['page_height_in']} in")
    print(f"Margins: {info['margins_in']}")
    print(f"Line numbers: {'YES' if info['line_numbers'] else 'NO'}")
    print()
    print("--- Typography ---")
    print(f"Fonts detected: {', '.join(info['fonts_detected'])}")
    print(f"Body paragraphs: {info['body_paragraphs']} ({info['double_spaced']} double-spaced)")
    print(f"Superscript runs: {info['superscript_runs']}")
    print(f"Greek letters: {', '.join(info['greek_letters']) if info['greek_letters'] else 'NONE'}")
    print(f"Special chars: {', '.join(info['special_characters']) if info['special_characters'] else 'NONE'}")
    print()
    print("--- Structure ---")
    print(f"Heading 1 sections: {info['heading_1_count']}")
    print(f"Heading 2 sections: {info['heading_2_count']}")
    print(f"Heading 3 sections: {info['heading_3_count']}")
    print(f"Tables: {info['table_count']}")
    print(f"Expected sections found: {info['expected_sections_found']}/{info['expected_sections_total']}")
    print(f"H1 sections: {info['h1_sections']}")
    print()
    if info['placeholders_preserved']:
        print(f"Placeholders preserved: {', '.join(info['placeholders_preserved'])}")
    print()

    if warnings:
        print(f"--- WARNINGS ({len(warnings)}) ---")
        for w in warnings:
            print(f"  ⚠ {w}")
    else:
        print("--- No warnings ---")
    print()
    print("=" * 60)

    return len(warnings)


if __name__ == "__main__":
    exit_code = main()
    sys.exit(min(exit_code, 1) if exit_code > 0 else 0)
