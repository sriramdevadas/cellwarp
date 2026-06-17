#!/usr/bin/env python3
# RETIRED: the Key Resources Table was removed in the PLOS reformat (D5 / WP-E.3). This script is no longer used and its output no longer exists.
"""
Build Key Resources Table Word document for Cell Systems submission.

Cell Systems requires the Key Resources Table as a separate Word document
uploaded in Editorial Manager.

Biology: Manuscript formatting — no new analysis.
Math: None.

Input:  docs/submission/key_resources_table.md
Output: docs/submission/key_resources_table.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "submission" / "key_resources_table.md"
OUTPUT = ROOT / "docs" / "submission" / "key_resources_table.docx"

FONT_NAME = "Arial"
TABLE_SIZE = Pt(9)
HEADING_SIZE = Pt(14)


def parse_table(lines: list) -> list:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        # Skip separator lines
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        if cells:
            rows.append(cells)
    return rows


def main():
    text = SOURCE.read_text(encoding="utf-8")

    # Strip markdown heading
    text = re.sub(r"^# Key Resources Table\s*\n+", "", text)

    # Parse table
    table_lines = [l for l in text.split("\n") if "|" in l or re.match(r"^[-:| ]+$", l.strip())]
    rows = parse_table(table_lines)

    if not rows or len(rows) < 2:
        raise ValueError("Could not parse Key Resources Table from source file")

    print(f"  Parsed {len(rows)} rows ({len(rows) - 1} data rows)")

    # Build document
    doc = Document()

    # Page setup: US Letter landscape, 1-inch margins
    for section in doc.sections:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = TABLE_SIZE

    # Heading
    p = doc.add_paragraph()
    run = p.add_run("Key Resources Table")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Build Word table
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = row.cells[j]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                is_bold = (i == 0)

                # Handle markdown bold markers
                cell_text_clean = cell_text
                if cell_text.startswith("**") and cell_text.endswith("**"):
                    cell_text_clean = cell_text[2:-2]
                    is_bold = True
                elif cell_text.startswith("**"):
                    cell_text_clean = cell_text.replace("**", "")
                    is_bold = True

                run = p.add_run(cell_text_clean)
                run.font.name = FONT_NAME
                run.font.size = TABLE_SIZE
                run.bold = is_bold

                pf = p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_after = Pt(2)
                pf.space_before = Pt(2)

    # Save
    doc.save(str(OUTPUT))
    print(f"\nSaved -> {OUTPUT}")


if __name__ == "__main__":
    main()
