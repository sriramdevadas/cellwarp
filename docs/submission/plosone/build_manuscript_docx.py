#!/usr/bin/env python3
"""
Build the PLOS ONE submission DOCX from the plosone manuscript source.

PLOS accepts DOC, DOCX or RTF, so docs/submission/plosone/manuscript_combined.txt
is not submittable as plain text, and the conversion must also deliver double
spacing, page numbers and continuous line numbers.

This script is written for one known input and refuses to guess. Every defect
found in this toolchain so far has been a silent one, so the rule here is that
a script which aborts is a good outcome and a script which emits plausible
wrong output is not:

  * the source is checked against reproduce/MANUSCRIPT_MD5 before it is
    parsed, so a submission DOCX cannot be built from an unpinned manuscript;
  * the nine section banners must match SECTIONS exactly, in order;
  * headings come from explicit line-number and text lists, never from a
    length or capitalization heuristic;
  * every content line must be claimed by exactly one handler, and the
    handler tallies must equal HANDLER_COUNTS;
  * the finished DOCX is re-opened and its extracted text must be
    character-identical to the source after the seven superscript
    substitutions, with the code-point inventory, the run formatting and the
    w:sectPr element order all asserted against the file actually written.

Usage:  python build_manuscript_docx.py [--output PATH]

Requires python-docx. As of 3f1d326 it is declared in the two extras the
documented installs actually read -- pyproject.toml's [lock], which is
`pip install -e ".[lock]"` at reproduce/README.md:35, and [dev], which is what
Dockerfile:65 installs and therefore what the container gets -- as well as in
requirements.txt, environment.yml and the [reproduce] extra. Before that it was
declared in [reproduce] alone, which no documented install reads, so this script
aborted at import in every environment the documentation tells you to build.

Note that the four gates run under .venv and this script does not: see
"Two interpreters" in reproduce/README.md.

The DOCX is a build artifact and is gitignored: a DOCX is a zip of XML carrying
creation timestamps, so it cannot be pinned to a stable hash the way the source
text is. No gate reads it, so a broken build here leaves all four gates green.
"""

import argparse
import hashlib
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - the dependency is declared
    sys.exit("ERROR: python-docx is required (pip install 'python-docx>=1.1'): %s" % exc)


# --- Paths -------------------------------------------------------------------

HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parent.parent.parent
SOURCE = HERE / "manuscript_combined.txt"
MD5_PIN = REPO_ROOT / "reproduce" / "MANUSCRIPT_MD5"
MD5_PIN_PATH = "docs/submission/plosone/manuscript_combined.txt"
DEFAULT_OUTPUT = HERE / "CellWarp_PLOSONE_manuscript.docx"


# --- Source structure --------------------------------------------------------

RULE = "=" * 72

# The nine section banners, in the order they must appear.
SECTIONS = (
    "TITLE",
    "Abstract",
    "INTRODUCTION",
    "RESULTS",
    "DISCUSSION",
    "Materials and Methods",
    "ACKNOWLEDGMENTS",
    "REFERENCES",
    "Supporting Information",
)

# 1-based source line numbers. The md5 gate pins the file, so a line number is
# as exact a reference as the text itself, and it is the only way to name
# Results heading 2 without retyping 104 characters.
TITLE_LINES = (6, 8, 10, 11, 12, 14)
RESULTS_HEADING_LINES = (47, 61, 77, 89, 101)
METHODS_HEADING_LINES = (132, 138, 146, 154, 160, 168, 172, 182, 192, 196)

# Belt and braces: the Methods sub-headings are short enough to state in full.
METHODS_HEADINGS = (
    "Ethics statement",
    "Data and cell-type matching",
    "The Procrustes framework",
    "The two-layer decomposition",
    "Primate replication (basal ganglia)",
    "Simulation",
    "Statistical analysis",
    "Conserved-contribution and identity-gene analysis",
    "Data and code availability",
    "Use of generative AI",
)

# The one body paragraph that legitimately has no terminal period. Any other
# body line without one is a sub-heading this script has not been told about.
BODY_WITHOUT_TERMINAL_PERIOD = (25,)

FIG_CAPTION = re.compile(r"^Fig \d+\. ")
SI_CAPTION = re.compile(r"^S\d+ (Fig|Table|Text)\. ")
NUMBERED = re.compile(r"^(\d+)\. ")

# A caption's label, title and legend go into one paragraph, with the title
# separated from its legend by this join. The round-trip gate is a character
# identity check, so the source side of the comparison inserts the same join at
# the same caption boundaries and nowhere else; reconstruct_source_lines() cuts
# every caption paragraph back apart to prove it. The joined counts below are a
# verification device rather than a requirement, so they move with the format.
CAPTION_JOIN = " "

HANDLER_COUNTS = {
    "title": 6,
    "h2": 15,
    "caption": 24,
    "legend": 24,
    "reference": 36,
    "body": 53,
}
EXPECTED_CONTENT_LINES = 158
EXPECTED_CAPTION_BOUNDARIES = 24  # 5 figure + 19 supporting information
# The 153 content lines joined with nothing between them: a property of the
# source alone, so it does not move when CAPTION_JOIN does.
EXPECTED_RAW_JOINED_CHARS = 94795
# The same join, plus CAPTION_JOIN at each of the 23 caption boundaries.
EXPECTED_JOINED_CHARS = 94819
EXPECTED_JOINED_WORDS = 13766
# ASCII T in the extracted text of the 121 content paragraphs: 253 in the source
# content lines plus 5 substituted from U+1D40. Content lines only. The eight
# emitted Heading 1 banners carry 4 more, and the TITLE banner's 2 never appear
# at all, because TITLE is rendered as a title page rather than as a heading.
EXPECTED_ASCII_T = 288
EXPECTED_REFERENCES = 36
EXPECTED_FIG_CAPTIONS = 5
EXPECTED_SI_CAPTIONS = 19


# --- Character-level formatting ----------------------------------------------

# Seven code points become real superscript runs. Every substitution is 1:1, so
# the character count is invariant and the round-trip gate can compare lengths.
# Escapes, not literals: U+207B and U+2212 are indistinguishable in source, and
# confusing them is exactly the kind of silent defect this script exists to
# prevent. The comment carries the Unicode name.
SUPERSCRIPT = {
    "\u00b9": "1",         # SUPERSCRIPT ONE
    "\u00b2": "2",         # SUPERSCRIPT TWO
    "\u2074": "4",         # SUPERSCRIPT FOUR
    "\u2075": "5",         # SUPERSCRIPT FIVE
    "\u2076": "6",         # SUPERSCRIPT SIX
    "\u207b": "\u2212",    # SUPERSCRIPT MINUS -> MINUS SIGN
    "\u1d40": "T",         # MODIFIER LETTER CAPITAL T -> transpose operator
}
SUPERSCRIPT_EXPECTED = {
    "\u00b9": 1,
    "\u00b2": 2,
    "\u2074": 32,
    "\u2075": 2,
    "\u2076": 9,
    "\u207b": 42,
    "\u1d40": 5,
}
SUPER_MINUS = "\u207b"
MINUS = "\u2212"

# Passed through unchanged, as literal characters in ordinary runs.
LITERAL_EXPECTED = {
    "\u00d7": 13,  # MULTIPLICATION SIGN
    "\u00e9": 1,   # LATIN SMALL LETTER E WITH ACUTE (Felix E, ChEMBL entry)
    "\u00f3": 1,   # LATIN SMALL LETTER O WITH ACUTE (Thorvaldsdottir H, MSigDB entry)
    "\u00f6": 1,   # LATIN SMALL LETTER O WITH DIAERESIS
    "\u0107": 1,   # LATIN SMALL LETTER C WITH ACUTE
    "\u03a3": 1,   # GREEK CAPITAL LETTER SIGMA
    "\u03b1": 1,   # GREEK SMALL LETTER ALPHA
    "\u03c1": 42,  # GREEK SMALL LETTER RHO
    "\u2013": 82,  # EN DASH
    "\u2014": 1,   # EM DASH
    "\u2016": 2,   # DOUBLE VERTICAL LINE
    "\u2032": 1,   # PRIME (was 3; two were the 3' v3 / 5' v2 sub-split, dropped)
    "\u2192": 8,   # RIGHTWARDS ARROW
    "\u2208": 4,   # ELEMENT OF
    "\u2212": 17,  # MINUS SIGN (59 in the output: 17 literal + 42 substituted)
    "\u2248": 11,  # ALMOST EQUAL TO
    "\u2264": 6,   # LESS-THAN OR EQUAL TO
    "\u2265": 5,   # GREATER-THAN OR EQUAL TO
}

# Species binomials to italicize, with the number of occurrences expected.
BINOMIALS = {"Homo sapiens": 1, "Microcebus murinus": 1,
             "Macaca mulatta": 3, "Macaca nemestrina": 1}

# Binomial-shaped pairs whose leading word is a genus but which are not species
# names, so they stay roman. Tabula Microcebus is the name of an atlas, and in
# "the Tabula Sapiens, Tabula Muris Senis, and Tabula Microcebus consortia" the
# genus word is followed by an English noun.
# "Macaca species" is the same shape: Macaca is a genus, but "species" is an
# English noun, not an epithet (the basal-ganglia macaque arm pools two species).
NOT_A_BINOMIAL = ("Microcebus consortia", "Macaca species")

# Coverage guard. Any "Capitalized lowercase" pair whose leading word is not
# classified below aborts the build, so a binomial introduced by a later edit
# cannot quietly render roman. The cost is that ordinary prose edits can also
# trip it; the fix is then a one-word addition here, and the error names the
# word and the line.
NOT_A_GENUS = (
    "Across", "After", "All", "An", "And", "At", "Atlas", "Automatic", "Base",
    "Because", "Benchmark", "Beyond", "Biological", "Bonferroni", "Bootstrap",
    "Both", "But", "Cell", "Cells", "Census", "Chromium", "Coherence",
    "Columns", "Combined", "Comparative", "Confidence", "Conserved",
    "Analysis", "Code", "Continuity", "Controlling", "Correlations",
    "Covariance", "More",
    "Corresponding", "Count", "Data", "Dated", "Statistical",
    "Dimensionality", "Dots", "Dropping", "Each", "Ellipsoid", "Ensembl",
    "Ethics", "Euclidean", "Every", "Evolutionarily", "Excluding", "Expanded",
    "Extended", "For", "Four", "Full", "Gaussian", "Gene", "Genes",
    "Geometric", "Global", "Hartigan", "Having", "Hochberg", "How", "Human",
    "In", "Intermediate", "It", "Italic", "Its", "Krzanowski", "Like",
    "Lower", "Mantel", "Mapping", "Marker", "Marmoset", "Materials",
    "Metazoa", "Molecular", "Most", "Myr", "Negative", "Neither", "No",
    "None", "Of", "Ontology", "Original", "Pearson", "Permutation",
    "Pipeline", "Plotting", "Position", "Precision", "Primary", "Primate",
    "Procrustes", "Progressive", "Protocol", "Python", "Random", "Rank",
    "Rankings", "Replacing", "Replication", "Reporting", "Representing",
    "Restricting", "Results", "Robustness", "Same", "Sapiens", "Scatter",
    "Senis", "Short", "Significance", "Simulation", "Software", "Source",
    "Spearman", "Splitting", "Substantiates", "Synthetic", "Table", "Taken",
    "Tau", "Ten", "Text", "That", "The", "This", "Three", "To", "Toward",
    "Treating", "Two", "Under", "Variance", "Wasserstein", "We", "Were",
    "What", "Whether", "Within", "Zenodo",
)
BINOMIAL_SHAPE = re.compile(r"\b[A-Z][a-z]+ [a-z]{3,}\b")


# --- Page setup --------------------------------------------------------------

FONT_NAME = "Times New Roman"  # PLOS prohibits only Symbol
FONT_SIZE = Pt(12)
HEADING1_SIZE = Pt(14)
TITLE_SIZE = Pt(16)
LINE_SPACING = 2.0
REFERENCE_HANGING_INDENT = Inches(0.5)

# w:lnNumType must precede w:cols and w:docGrid: CT_SectPr's schema fixes the
# child order, and sectPr.append() would place it after both, where Word drops
# it. python-docx 1.2.0 exposes no _tag_seq for CT_SectPr, so the successor
# list is spelled out here from the CT_SectPr sequence in the OOXML schema.
LNNUMTYPE_SUCCESSORS = (
    "w:pgNumType", "w:cols", "w:formProt", "w:vAlign", "w:noEndnote",
    "w:titlePg", "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
    "w:printerSettings", "w:sectPrChange",
)
COLS_SUCCESSORS = LNNUMTYPE_SUCCESSORS[2:]

# Word prefers a theme font over an explicit one, so the theme attributes have
# to go before FONT_NAME takes effect.
THEME_FONT_ATTRS = ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme")
EXPLICIT_FONT_ATTRS = ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia")


# --- Failure -----------------------------------------------------------------

class BuildError(Exception):
    """Raised for any violated structural assumption. Always fatal."""


def require(condition, message):
    """assert, but not stripped by python -O."""
    if not condition:
        raise BuildError(message)


# --- Source loading ----------------------------------------------------------

def read_pinned_md5():
    require(MD5_PIN.is_file(), "md5 pin file missing: %s" % MD5_PIN)
    matches = []
    for line in MD5_PIN.read_text(encoding="utf-8").splitlines():
        parts = line.split()
        if len(parts) == 2 and parts[1].lstrip("*") == MD5_PIN_PATH:
            matches.append(parts[0])
    require(
        len(matches) == 1,
        "expected exactly one entry for %s in %s, found %d"
        % (MD5_PIN_PATH, MD5_PIN, len(matches)),
    )
    return matches[0]


def load_source():
    """Return the source lines, refusing to proceed unless the md5 pin matches."""
    require(SOURCE.is_file(), "manuscript source missing: %s" % SOURCE)
    raw = SOURCE.read_bytes()
    digest = hashlib.md5(raw).hexdigest()
    pinned = read_pinned_md5()
    require(
        digest == pinned,
        "manuscript md5 mismatch: %s is %s but %s pins %s. Refusing to build a "
        "submission DOCX from an unpinned manuscript."
        % (SOURCE.name, digest, MD5_PIN.name, pinned),
    )
    lines = raw.decode("utf-8").split("\n")
    if lines and lines[-1] == "":
        lines.pop()
    return lines, digest


# --- Parsing -----------------------------------------------------------------

def split_sections(lines):
    """Return (section_of, content_indices) after validating the banner block."""
    rules = {i for i, line in enumerate(lines) if line == RULE}
    stray = [i + 1 for i, line in enumerate(lines)
             if line and set(line) == {"="} and line != RULE]
    require(not stray, "lines of '=' that are not exactly %d characters: %s"
            % (len(RULE), stray))
    require(len(rules) == 2 * len(SECTIONS),
            "expected %d rule lines, found %d" % (2 * len(SECTIONS), len(rules)))

    banners = sorted(i + 1 for i in rules if i + 2 in rules)
    names = [lines[i] for i in banners]
    require(
        tuple(names) == SECTIONS,
        "section banners do not match the expected nine, in order.\n"
        "  expected: %s\n  found:    %s" % (list(SECTIONS), names),
    )

    section_of = {}
    current = None
    for i, line in enumerate(lines):
        if i in banners:
            current = line
        section_of[i] = current

    banner_set = set(banners)
    content = [i for i, line in enumerate(lines)
               if line.strip() and i not in rules and i not in banner_set]
    require(
        len(content) == EXPECTED_CONTENT_LINES,
        "expected %d content lines, found %d"
        % (EXPECTED_CONTENT_LINES, len(content)),
    )
    ragged = [i + 1 for i in content if lines[i] != lines[i].strip()]
    require(not ragged, "content lines with leading or trailing whitespace: %s" % ragged)
    tabbed = [i + 1 for i in content if "\t" in lines[i]]
    require(not tabbed,
            "content lines containing a TAB (tables are not supported): %s" % tabbed)
    return section_of, content


def classify(lines, section_of, content):
    """Assign every content line to exactly one handler.

    Returns (plan, handler), where plan is the emit order as a list of
    (kind, text, label_length, source_line_numbers).
    """
    handler = {}
    plan = []
    heading_emitted = None
    order = {index: position for position, index in enumerate(content)}
    consumed = set()

    def claim(index, kind):
        require(index + 1 not in handler,
                "line %d claimed twice (%s then %s)"
                % (index + 1, handler.get(index + 1), kind))
        handler[index + 1] = kind

    for index in content:
        lineno = index + 1
        if lineno in consumed:
            continue
        text = lines[index]
        section = section_of[index]

        if section != SECTIONS[0] and heading_emitted != section:
            plan.append(("h1", section, 0, ()))
            heading_emitted = section

        if section == SECTIONS[0]:
            require(lineno in TITLE_LINES,
                    "unexpected content line %d in the TITLE section" % lineno)
            claim(index, "title")
            plan.append(("title", text, 0, (lineno,)))
            continue

        if lineno in RESULTS_HEADING_LINES or lineno in METHODS_HEADING_LINES:
            claim(index, "h2")
            plan.append(("h2", text, 0, (lineno,)))
            continue

        match = FIG_CAPTION.match(text) or SI_CAPTION.match(text)
        if match:
            position = order[index]
            require(position + 1 < len(content),
                    "caption at line %d has no following legend line" % lineno)
            legend_index = content[position + 1]
            legend = lines[legend_index]
            require(section_of[legend_index] == section,
                    "legend for the caption at line %d crosses a section boundary"
                    % lineno)
            require(not (FIG_CAPTION.match(legend) or SI_CAPTION.match(legend)),
                    "caption at line %d is followed by another caption, not a legend"
                    % lineno)
            require(legend_index + 1 not in RESULTS_HEADING_LINES
                    and legend_index + 1 not in METHODS_HEADING_LINES,
                    "caption at line %d is followed by a heading, not a legend" % lineno)
            claim(index, "caption")
            claim(legend_index, "legend")
            consumed.add(legend_index + 1)
            plan.append(("caption", text + CAPTION_JOIN + legend, match.end(),
                         (lineno, legend_index + 1)))
            continue

        if section == "REFERENCES":
            require(NUMBERED.match(text),
                    "line %d in REFERENCES is not a numbered reference" % lineno)
            claim(index, "reference")
            plan.append(("reference", text, 0, (lineno,)))
            continue

        require(
            text.endswith(".") or lineno in BODY_WITHOUT_TERMINAL_PERIOD,
            "line %d has no terminal period and is not a known heading or a "
            "declared exception. If it is a new sub-heading, add it to the "
            "heading lists; if it is prose, add it to "
            "BODY_WITHOUT_TERMINAL_PERIOD.\n  %r" % (lineno, text),
        )
        require(
            not NUMBERED.match(text),
            "line %d starts with a number and a period but is neither a listed "
            "heading nor a reference:\n  %r" % (lineno, text),
        )
        claim(index, "body")
        plan.append(("body", text, 0, (lineno,)))

    tally = Counter(handler.values())
    require(dict(tally) == HANDLER_COUNTS,
            "handler tally mismatch.\n  expected: %s\n  found:    %s"
            % (HANDLER_COUNTS, dict(tally)))
    require(len(handler) == EXPECTED_CONTENT_LINES,
            "handlers claimed %d lines, expected %d"
            % (len(handler), EXPECTED_CONTENT_LINES))

    validate_headings(section_of, plan)
    validate_captions(plan)
    validate_references(plan)
    validate_species(plan)
    return plan, handler


def validate_headings(section_of, plan):
    results = [p for p in plan if p[0] == "h2" and p[3][0] in RESULTS_HEADING_LINES]
    methods = [p for p in plan if p[0] == "h2" and p[3][0] in METHODS_HEADING_LINES]
    require(len(results) == len(RESULTS_HEADING_LINES),
            "expected %d Results sub-headings, found %d"
            % (len(RESULTS_HEADING_LINES), len(results)))
    require(len(methods) == len(METHODS_HEADING_LINES),
            "expected %d Methods sub-headings, found %d"
            % (len(METHODS_HEADING_LINES), len(methods)))

    numbers = []
    for _, text, _, source in results:
        require(section_of[source[0] - 1] == "RESULTS",
                "Results sub-heading at line %d is not in RESULTS" % source[0])
        match = NUMBERED.match(text)
        require(match, "Results sub-heading at line %d is not numbered:\n  %r"
                % (source[0], text))
        numbers.append(int(match.group(1)))
    require(numbers == list(range(1, len(RESULTS_HEADING_LINES) + 1)),
            "Results sub-headings are not numbered 1..%d in order: %s"
            % (len(RESULTS_HEADING_LINES), numbers))

    found = tuple(text for _, text, _, _ in methods)
    require(found == METHODS_HEADINGS,
            "Methods sub-headings do not match the expected list.\n"
            "  expected: %s\n  found:    %s" % (list(METHODS_HEADINGS), list(found)))
    for _, _, _, source in methods:
        require(section_of[source[0] - 1] == "Materials and Methods",
                "Methods sub-heading at line %d is not in Materials and Methods"
                % source[0])

    for _, text, _, source in results + methods:
        require(not text.endswith("."),
                "sub-heading at line %d ends with a period:\n  %r" % (source[0], text))


def validate_captions(plan):
    figs, sis = [], []
    for kind, text, label_length, source in plan:
        if kind != "caption":
            continue
        label = text[:label_length]
        if FIG_CAPTION.fullmatch(label):
            figs.append(source[0])
        elif SI_CAPTION.fullmatch(label):
            sis.append(source[0])
        else:
            raise BuildError("caption at line %d has an unrecognized label: %r"
                             % (source[0], label))
        require(label.endswith(". "),
                "caption label at line %d does not end with '. ': %r"
                % (source[0], label))
    require(len(figs) == EXPECTED_FIG_CAPTIONS,
            "expected %d figure captions, found %d"
            % (EXPECTED_FIG_CAPTIONS, len(figs)))
    require(len(sis) == EXPECTED_SI_CAPTIONS,
            "expected %d supporting-information captions, found %d"
            % (EXPECTED_SI_CAPTIONS, len(sis)))
    require(len(figs) + len(sis) == EXPECTED_CAPTION_BOUNDARIES,
            "expected %d caption boundaries to receive CAPTION_JOIN, found %d"
            % (EXPECTED_CAPTION_BOUNDARIES, len(figs) + len(sis)))


def validate_references(plan):
    numbers = [int(NUMBERED.match(text).group(1))
               for kind, text, _, _ in plan if kind == "reference"]
    require(len(numbers) == EXPECTED_REFERENCES,
            "expected %d references, found %d" % (EXPECTED_REFERENCES, len(numbers)))
    require(numbers == list(range(1, EXPECTED_REFERENCES + 1)),
            "references are not numbered 1..%d in order: %s"
            % (EXPECTED_REFERENCES, numbers))


def validate_species(plan):
    """Every binomial must be handled, and every binomial-shaped pair classified."""
    counts = Counter()
    unclassified = []
    for _, text, _, source in plan:
        for name in BINOMIALS:
            counts[name] += text.count(name)
        for match in BINOMIAL_SHAPE.finditer(text):
            pair = match.group(0)
            if pair.split(" ", 1)[0] in NOT_A_GENUS or pair in NOT_A_BINOMIAL:
                continue
            if pair not in BINOMIALS:
                unclassified.append((source[0], pair))
    require(
        not unclassified,
        "binomial-shaped pairs that are neither a handled species name nor a "
        "classified non-genus word. Add the species to BINOMIALS or its leading "
        "word to NOT_A_GENUS:\n  %s"
        % "\n  ".join("line %d: %r" % item for item in unclassified),
    )
    for name, expected in BINOMIALS.items():
        require(counts[name] == expected,
                "expected %d occurrence(s) of %r, found %d"
                % (expected, name, counts[name]))


# --- Run building ------------------------------------------------------------

def build_spans(text, label_length=0):
    """Split text into (text, bold, italic, superscript) runs.

    Substitution is 1:1 per character, so offsets into the original text stay
    valid for the italic spans.
    """
    chars = list(text)
    bold = [i < label_length for i in range(len(chars))]
    italic = [False] * len(chars)
    superscript = [False] * len(chars)

    for i, char in enumerate(chars):
        if char in SUPERSCRIPT:
            chars[i] = SUPERSCRIPT[char]
            superscript[i] = True

    for name in BINOMIALS:
        start = text.find(name)
        while start >= 0:
            for i in range(start, start + len(name)):
                italic[i] = True
            start = text.find(name, start + len(name))

    spans = []
    for i, char in enumerate(chars):
        key = (bold[i], italic[i], superscript[i])
        if spans and spans[-1][0] == key:
            spans[-1][1].append(char)
        else:
            spans.append((key, [char]))
    return [("".join(buffer), key[0], key[1], key[2]) for key, buffer in spans]


def add_runs(paragraph, text, label_length=0):
    for content, bold, italic, superscript in build_spans(text, label_length):
        run = paragraph.add_run(content)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if superscript:
            run.font.superscript = True
    return paragraph


# --- Document assembly -------------------------------------------------------

def set_style_font(style, size, bold=None):
    style.font.name = FONT_NAME
    style.font.size = size
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if bold is not None:
        style.font.bold = bold
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in THEME_FONT_ATTRS:
        if fonts.get(qn(attr)) is not None:
            del fonts.attrib[qn(attr)]
    for attr in EXPLICIT_FONT_ATTRS:
        fonts.set(qn(attr), FONT_NAME)


def configure_styles(document):
    set_style_font(document.styles["Normal"], FONT_SIZE)
    for name, size in (("Heading 1", HEADING1_SIZE), ("Heading 2", FONT_SIZE)):
        style = document.styles[name]
        set_style_font(style, size, bold=True)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    """PAGE field, as the begin / instruction / end run triple Word expects."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        paragraph.add_run()._r.append(element)


def configure_section(section):
    """Single column, page-number footer, continuous line numbers."""
    sect_pr = section._sectPr

    section.footer.is_linked_to_previous = False
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_paragraph)

    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.insert_element_before(cols, *COLS_SUCCESSORS)
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "720")

    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.insert_element_before(line_numbers, *LNNUMTYPE_SUCCESSORS)


def strip_package_thumbnail(document):
    """Drop docProps/thumbnail.jpeg and its package relationship.

    python-docx's default template ships a thumbnail of a blank page, which then
    rides along in a file going to a journal. Deleting the zip entry alone would
    leave a dangling relationship in _rels/.rels and an invalid package, so the
    relationship is what gets removed: the serializer reaches parts by walking
    the relationship graph, so the part goes with it and the jpeg default drops
    out of [Content_Types].xml.
    """
    rels = document.part.package.rels
    dropped = [rid for rid, rel in rels.items()
               if rel.reltype == RELATIONSHIP_TYPE.THUMBNAIL]
    for rid in dropped:
        del rels[rid]
    return dropped


def build_document(plan):
    document = Document()
    strip_package_thumbnail(document)
    configure_styles(document)
    configure_section(document.sections[0])

    title_done = False
    break_done = False
    for kind, text, label_length, _ in plan:
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(paragraph, text)
            if not title_done:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = TITLE_SIZE
                title_done = True
        elif kind == "h1":
            paragraph = document.add_paragraph(style="Heading 1")
            add_runs(paragraph, text)
            if not break_done:
                paragraph.paragraph_format.page_break_before = True
                break_done = True
        elif kind == "h2":
            paragraph = document.add_paragraph(style="Heading 2")
            add_runs(paragraph, text)
        elif kind == "reference":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = REFERENCE_HANGING_INDENT
            paragraph.paragraph_format.first_line_indent = -REFERENCE_HANGING_INDENT
            add_runs(paragraph, text)
        else:
            paragraph = document.add_paragraph()
            add_runs(paragraph, text, label_length)
        paragraph.paragraph_format.line_spacing = LINE_SPACING
    return document


# --- Verification of the file actually written -------------------------------

def substitute(text):
    for source, target in SUPERSCRIPT.items():
        text = text.replace(source, target)
    return text


def local_name(tag):
    return tag.rpartition("}")[2]


def reconstruct_source_lines(content_texts, plan, lines, content):
    """Cut every caption paragraph back apart at its known boundary.

    The paragraph-level check proves each paragraph matches the plan. This
    proves the plan matches the source: splitting the 23 caption paragraphs at
    the length of their caption line has to return the 144 source lines exactly,
    which is only true if CAPTION_JOIN went in at those 23 boundaries and at no
    other point in the document.
    """
    rebuilt = []
    boundaries = 0
    for text, (kind, _, _, source) in zip(
            content_texts, [entry for entry in plan if entry[0] != "h1"]):
        if kind != "caption":
            rebuilt.append(text)
            continue
        head = substitute(lines[source[0] - 1])
        tail = substitute(lines[source[1] - 1])
        cut = len(head)
        require(text[:cut] == head,
                "caption paragraph from line %d does not start with its caption "
                "line:\n  expected: %r\n  found:    %r" % (source[0], head, text[:cut]))
        require(text[cut:cut + len(CAPTION_JOIN)] == CAPTION_JOIN,
                "caption from line %d is not followed by CAPTION_JOIN %r, but by %r"
                % (source[0], CAPTION_JOIN, text[cut:cut + len(CAPTION_JOIN)]))
        require(text[cut + len(CAPTION_JOIN):] == tail,
                "legend from line %d does not follow the join:\n"
                "  expected: %r\n  found:    %r"
                % (source[1], tail, text[cut + len(CAPTION_JOIN):]))
        rebuilt.extend([head, tail])
        boundaries += 1
    require(boundaries == EXPECTED_CAPTION_BOUNDARIES,
            "CAPTION_JOIN was inserted at %d boundaries, expected %d"
            % (boundaries, EXPECTED_CAPTION_BOUNDARIES))
    expected = [substitute(lines[index]) for index in content]
    require(rebuilt == expected,
            "the document does not cut back apart into the %d source lines "
            "(%d reconstructed)" % (len(expected), len(rebuilt)))
    return boundaries


def verify(path, plan, lines, content):
    """Re-open the written DOCX and check it against the source."""
    raw_joined = "".join(lines[index] for index in content)
    require(len(raw_joined) == EXPECTED_RAW_JOINED_CHARS,
            "the %d source content lines join to %d characters, expected %d"
            % (len(content), len(raw_joined), EXPECTED_RAW_JOINED_CHARS))
    source_joined = "".join(text for kind, text, _, _ in plan if kind != "h1")
    require(len(source_joined)
            == EXPECTED_RAW_JOINED_CHARS
            + EXPECTED_CAPTION_BOUNDARIES * len(CAPTION_JOIN),
            "the source side joins to %d characters, expected %d plus %d joins of %r"
            % (len(source_joined), EXPECTED_RAW_JOINED_CHARS,
               EXPECTED_CAPTION_BOUNDARIES, CAPTION_JOIN))

    document = Document(str(path))
    paragraphs = document.paragraphs
    styles = [p.style.name for p in paragraphs]
    texts = [p.text for p in paragraphs]

    empty = [i for i, text in enumerate(texts) if text == ""]
    require(not empty, "the document contains %d empty paragraph(s), at %s"
            % (len(empty), empty))

    heading1 = [t for t, s in zip(texts, styles) if s == "Heading 1"]
    heading2 = [t for t, s in zip(texts, styles) if s == "Heading 2"]
    deeper = sorted({s for s in styles
                     if s.startswith("Heading ") and s not in ("Heading 1", "Heading 2")})
    require(not deeper, "heading depth exceeds 2: %s" % deeper)
    require(tuple(heading1) == SECTIONS[1:],
            "Heading 1 paragraphs do not match the eight emitted banners.\n"
            "  expected: %s\n  found:    %s" % (list(SECTIONS[1:]), heading1))
    require(len(heading2) == HANDLER_COUNTS["h2"],
            "expected %d Heading 2 paragraphs, found %d"
            % (HANDLER_COUNTS["h2"], len(heading2)))

    expected_texts = [substitute(text) for kind, text, _, _ in plan if kind != "h1"]
    content_texts = [t for t, s in zip(texts, styles) if s != "Heading 1"]
    require(len(content_texts) == len(expected_texts),
            "expected %d content paragraphs, found %d"
            % (len(expected_texts), len(content_texts)))
    for i, (found, expected) in enumerate(zip(content_texts, expected_texts)):
        require(found == expected,
                "content paragraph %d differs from the source:\n"
                "  expected: %r\n  found:    %r" % (i, expected, found))

    joined = "".join(content_texts)
    expected_joined = substitute(source_joined)
    require(joined == expected_joined,
            "round trip is not character-identical (%d vs %d characters)"
            % (len(joined), len(expected_joined)))
    require(len(joined) == EXPECTED_JOINED_CHARS,
            "joined character count is %d, expected %d"
            % (len(joined), EXPECTED_JOINED_CHARS))
    require(len(joined.split()) == EXPECTED_JOINED_WORDS,
            "joined word count is %d, expected %d"
            % (len(joined.split()), EXPECTED_JOINED_WORDS))
    boundaries = reconstruct_source_lines(content_texts, plan, lines, content)

    require(joined.count("T") == EXPECTED_ASCII_T,
            "ASCII T count in the content paragraphs is %d, expected %d"
            % (joined.count("T"), EXPECTED_ASCII_T))

    for char in SUPERSCRIPT:
        require(joined.count(char) == 0,
                "U+%04X survived into the extracted text (%d occurrences)"
                % (ord(char), joined.count(char)))
    expected_minus = LITERAL_EXPECTED[MINUS] + SUPERSCRIPT_EXPECTED[SUPER_MINUS]
    require(joined.count(MINUS) == expected_minus,
            "U+2212 count is %d, expected %d" % (joined.count(MINUS), expected_minus))
    for char, count in LITERAL_EXPECTED.items():
        if char == MINUS:
            continue
        require(joined.count(char) == count,
                "U+%04X count is %d, expected %d"
                % (ord(char), joined.count(char), count))

    require(not document.tables,
            "the document contains %d table(s)" % len(document.tables))
    require(not document.inline_shapes,
            "the document contains %d inline shape(s)" % len(document.inline_shapes))

    runs = [r for p in paragraphs for r in p.runs]
    bold_runs = [r.text for r in runs if r.bold]
    italic_runs = [r.text for r in runs if r.italic]
    superscript_runs = [r.text for r in runs if r.font.superscript]
    fig_labels = [t for t in bold_runs if FIG_CAPTION.fullmatch(t)]
    si_labels = [t for t in bold_runs if SI_CAPTION.fullmatch(t)]
    require(len(fig_labels) == EXPECTED_FIG_CAPTIONS,
            "expected %d bold figure labels, found %d"
            % (EXPECTED_FIG_CAPTIONS, len(fig_labels)))
    require(len(si_labels) == EXPECTED_SI_CAPTIONS,
            "expected %d bold supporting-information labels, found %d"
            % (EXPECTED_SI_CAPTIONS, len(si_labels)))
    # One italic run per OCCURRENCE, so compare against BINOMIALS expanded by its
    # declared counts. Comparing against sorted(BINOMIALS) was correct only while
    # every binomial occurred exactly once.
    expected_italics = sorted(name for name, n in BINOMIALS.items() for _ in range(n))
    require(sorted(italic_runs) == expected_italics,
            "italic runs are %s, expected %s"
            % (sorted(italic_runs), expected_italics))
    superscript_chars = sum(len(t) for t in superscript_runs)
    require(superscript_chars == sum(SUPERSCRIPT_EXPECTED.values()),
            "superscript runs carry %d characters, expected %d"
            % (superscript_chars, sum(SUPERSCRIPT_EXPECTED.values())))

    return {
        "paragraphs": len(paragraphs),
        "content_paragraphs": len(content_texts),
        "heading1": len(heading1),
        "heading2": len(heading2),
        "joined_chars": len(joined),
        "joined_words": len(joined.split()),
        "identical": joined == expected_joined,
        "joined_with_headings": len("".join(texts)),
        "ascii_t_content": joined.count("T"),
        "ascii_t_all": "".join(texts).count("T"),
        "minus": joined.count(MINUS),
        "literals": {char: joined.count(char) for char in LITERAL_EXPECTED},
        "bold_runs": len(bold_runs),
        "fig_labels": len(fig_labels),
        "si_labels": len(si_labels),
        "italic_runs": italic_runs,
        "superscript_runs": len(superscript_runs),
        "superscript_chars": superscript_chars,
        "caption_boundaries": boundaries,
        "raw_joined_chars": len(raw_joined),
        "xml": verify_xml(path, len(paragraphs), len(superscript_runs)),
    }


def verify_xml(path, paragraph_count, superscript_run_count):
    """Assert the structural facts only the written XML can settle."""
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        footers = [n for n in names if re.fullmatch(r"word/footer\d*\.xml", n)]
        require(footers, "no footer part in the package: %s" % names)
        footer_xml = "".join(archive.read(n).decode("utf-8") for n in footers)
        package_rels = archive.read("_rels/.rels").decode("utf-8")
        content_types = archive.read("[Content_Types].xml").decode("utf-8")

    thumbnails = [n for n in names if "thumbnail" in n.lower()]
    require(not thumbnails, "the package still carries a thumbnail part: %s" % thumbnails)
    require("thumbnail" not in package_rels.lower(),
            "_rels/.rels still references a thumbnail, which would leave the "
            "package invalid once the part is gone")
    require("jpeg" not in content_types.lower(),
            "[Content_Types].xml still declares a jpeg, but no image part remains")

    require("w:footnoteReference" not in document_xml,
            "the document contains footnote references")
    require("<w:drawing" not in document_xml and "<w:pict" not in document_xml,
            "the document contains embedded graphics")
    require("<w:tbl>" not in document_xml, "the document contains a table")
    require(re.search(r"<w:instrText[^>]*>\s*PAGE\s*</w:instrText>", footer_xml),
            "no PAGE field in the footer")
    require("w:footerReference" in document_xml,
            "the section does not reference a footer")

    sect_pr = Document(str(path)).sections[0]._sectPr
    order = [local_name(child.tag) for child in sect_pr]
    require("lnNumType" in order, "w:lnNumType is absent from w:sectPr: %s" % order)
    for successor in ("cols", "docGrid"):
        require(successor in order,
                "w:%s is absent from w:sectPr: %s" % (successor, order))
        require(order.index("lnNumType") < order.index(successor),
                "w:lnNumType is at position %d, after w:%s at position %d, which "
                "violates the CT_SectPr sequence: %s"
                % (order.index("lnNumType"), successor, order.index(successor), order))
    line_numbers = sect_pr.find(qn("w:lnNumType"))
    require(line_numbers.get(qn("w:countBy")) == "1",
            "w:lnNumType/@w:countBy is %r" % line_numbers.get(qn("w:countBy")))
    require(line_numbers.get(qn("w:restart")) == "continuous",
            "w:lnNumType/@w:restart is %r" % line_numbers.get(qn("w:restart")))
    cols = sect_pr.find(qn("w:cols"))
    require(cols.get(qn("w:num")) == "1",
            "w:cols/@w:num is %r, expected a single column" % cols.get(qn("w:num")))

    spacing_tags = re.findall(r"<w:spacing\b[^>]*>", document_xml)
    spaced = sum(1 for tag in spacing_tags
                 if 'w:line="480"' in tag and 'w:lineRule="auto"' in tag)
    require(spaced == paragraph_count,
            "%d of %d paragraphs carry w:line=480 with w:lineRule=auto"
            % (spaced, paragraph_count))
    vert_align = len(re.findall(r'<w:vertAlign w:val="superscript"\s*/>', document_xml))
    require(vert_align == superscript_run_count,
            "%d w:vertAlign superscript elements for %d superscript runs"
            % (vert_align, superscript_run_count))

    return {
        "sect_pr_order": order,
        "ln_num_type_index": order.index("lnNumType"),
        "cols_index": order.index("cols"),
        "doc_grid_index": order.index("docGrid"),
        "count_by": line_numbers.get(qn("w:countBy")),
        "restart": line_numbers.get(qn("w:restart")),
        "columns": cols.get(qn("w:num")),
        "double_spaced": spaced,
        "vert_align": vert_align,
        "parts": len(names),
        "thumbnail_parts": len(thumbnails),
        "thumbnail_rels": package_rels.lower().count("thumbnail"),
    }


# --- Report ------------------------------------------------------------------

def print_report(path, digest, handler, report):
    xml = report["xml"]
    tally = Counter(handler.values())
    print("SOURCE  %s" % SOURCE)
    print("  md5 matches %-24s %s" % (MD5_PIN.name, digest))
    print("  section banners                     %d" % len(SECTIONS))
    print("  content lines                       %d" % len(handler))
    for kind in sorted(HANDLER_COUNTS):
        print("    %-33s %d" % (kind, tally[kind]))
    print("  figure / SI captions                %d / %d"
          % (EXPECTED_FIG_CAPTIONS, EXPECTED_SI_CAPTIONS))
    print("  references                          %d" % EXPECTED_REFERENCES)
    print("  binomials italicized                %s" % ", ".join(sorted(BINOMIALS)))
    print()
    print("OUTPUT  %s" % path)
    print("  bytes                               %d" % path.stat().st_size)
    print("  font                                %s %gpt, single column, %g line spacing"
          % (FONT_NAME, FONT_SIZE.pt, LINE_SPACING))
    print("  paragraphs, all / content           %d / %d"
          % (report["paragraphs"], report["content_paragraphs"]))
    print("  Heading 1 / Heading 2               %d / %d"
          % (report["heading1"], report["heading2"]))
    print()
    print("ROUND TRIP")
    print("  source content lines                %d" % EXPECTED_CONTENT_LINES)
    print("  source lines joined bare            %d" % report["raw_joined_chars"])
    print("  caption boundaries taking %-9r %d"
          % (CAPTION_JOIN, report["caption_boundaries"]))
    print("  joined characters                   %d" % report["joined_chars"])
    print("  joined words                        %d" % report["joined_words"])
    print("  character-identical                 %s" % report["identical"])
    print("  cuts back apart into source lines   %d" % EXPECTED_CONTENT_LINES)
    print("  joined incl. Heading 1 paragraphs   %d" % report["joined_with_headings"])
    print()
    print("CODE POINTS in the extracted text")
    for char in sorted(SUPERSCRIPT):
        print("  U+%04X substituted away              %d"
              % (ord(char), 0))
    # Derived, not spelled out: this line read "15 literal + 41 sub" while the
    # constants said 16 and 42, and it had been wrong for long enough that the
    # numbers no longer resembled anything. It cannot go stale again.
    print("  U+2212 minus, %d literal + %d sub    %d"
          % (LITERAL_EXPECTED[MINUS], SUPERSCRIPT_EXPECTED[SUPER_MINUS], report["minus"]))
    for char in sorted(LITERAL_EXPECTED):
        if char == MINUS:
            continue
        print("  U+%04X literal                       %d  (expected %d)"
              % (ord(char), report["literals"][char], LITERAL_EXPECTED[char]))
    print("  ASCII T, content paragraphs         %d" % report["ascii_t_content"])
    print("  ASCII T, incl. Heading 1            %d" % report["ascii_t_all"])
    print()
    print("XML  word/document.xml")
    print("  w:sectPr children                   %s" % " ".join(xml["sect_pr_order"]))
    print("  lnNumType / cols / docGrid index    %d / %d / %d"
          % (xml["ln_num_type_index"], xml["cols_index"], xml["doc_grid_index"]))
    print("  lnNumType countBy / restart         %s / %s"
          % (xml["count_by"], xml["restart"]))
    print("  columns                             %s" % xml["columns"])
    print("  paragraphs at w:line=480 auto       %d" % xml["double_spaced"])
    print("  PAGE field in footer                yes")
    print("  w:vertAlign superscript elements    %d, carrying %d characters"
          % (xml["vert_align"], report["superscript_chars"]))
    print("  bold runs, total / Fig / S          %d / %d / %d"
          % (report["bold_runs"], report["fig_labels"], report["si_labels"]))
    print("  italic runs                         %d  %s"
          % (len(report["italic_runs"]), ", ".join(report["italic_runs"])))
    print("  package parts                       %d" % xml["parts"])
    print("  thumbnail parts / relationships     %d / %d"
          % (xml["thumbnail_parts"], xml["thumbnail_rels"]))
    print()
    print("OK  %s" % path)


# --- Entry point -------------------------------------------------------------

def parse_args(argv):
    parser = argparse.ArgumentParser(
        description="Build the PLOS ONE submission DOCX from manuscript_combined.txt.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT,
                        help="where to write the DOCX (default: %(default)s)")
    return parser.parse_args(argv)


def check_constants():
    """The joined constants have to stay consistent with the caption format."""
    require(EXPECTED_JOINED_CHARS
            == EXPECTED_RAW_JOINED_CHARS
            + EXPECTED_CAPTION_BOUNDARIES * len(CAPTION_JOIN),
            "EXPECTED_JOINED_CHARS is %d, but %d source characters plus %d joins "
            "of %r is %d. Changing CAPTION_JOIN moves this constant, and "
            "EXPECTED_JOINED_WORDS with it."
            % (EXPECTED_JOINED_CHARS, EXPECTED_RAW_JOINED_CHARS,
               EXPECTED_CAPTION_BOUNDARIES, CAPTION_JOIN,
               EXPECTED_RAW_JOINED_CHARS
               + EXPECTED_CAPTION_BOUNDARIES * len(CAPTION_JOIN)))
    require(EXPECTED_CAPTION_BOUNDARIES == EXPECTED_FIG_CAPTIONS + EXPECTED_SI_CAPTIONS,
            "EXPECTED_CAPTION_BOUNDARIES is %d, but there are %d figure and %d "
            "supporting-information captions"
            % (EXPECTED_CAPTION_BOUNDARIES, EXPECTED_FIG_CAPTIONS, EXPECTED_SI_CAPTIONS))


def main(argv=None):
    args = parse_args(argv)
    try:
        check_constants()
        lines, digest = load_source()
        section_of, content = split_sections(lines)
        plan, handler = classify(lines, section_of, content)
        document = build_document(plan)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(args.output))
        report = verify(args.output, plan, lines, content)
    except BuildError as error:
        sys.exit("ERROR: %s" % error)
    print_report(args.output, digest, handler, report)
    return 0


if __name__ == "__main__":
    sys.exit(main())
